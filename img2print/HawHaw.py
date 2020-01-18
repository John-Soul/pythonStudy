import os, datetime, shutil
from os import listdir
from PIL import Image 
from docx import Document
from docx.shared import Inches

myDoc = Document()
today = str(datetime.date.today())
localDir = os.path.dirname(os.path.abspath(__file__))
historyImgsDir = os.path.dirname(os.path.abspath(__file__)) + '\\history_imgs\\' + today + '磅单原始照片'
pictures = []
isFirstImg = True

# 新建文件夹存储已经用过的磅单图片
def mkdir(path):
    path=path.strip()
    path=path.rstrip("\\")
    isExists=os.path.exists(path)
 
    if not isExists:
        os.makedirs(path) 
        return True
    else:
        return False

mkdir(historyImgsDir)

for item in listdir(localDir + '\imgs'):
    fileType = os.path.splitext(item)[1]
    fileName = os.path.splitext(item)[0]
    if fileType == '.jpg' or fileType == '.png' or fileType ==  '.JPG' or fileType == '.jpeg':
        pictures.append({
            'path': localDir + '\\imgs\\' + item,
            'fileName': fileName
        })

# print(pictures)
# pictures.sort()
# print(pictures)


for fn in pictures:
    img = Image.open(fn['path'])
    height = img.height
    width = img.width
    if (height > width):
        out = img.transpose(Image.ROTATE_90)
        out.save(fn['path'])
    myDoc.add_paragraph(fn['fileName'])
    myDoc.add_picture(fn['path'], width=Inches(6), height=Inches(3.15))
    if(isFirstImg):
        myDoc.add_paragraph()
        myDoc.add_paragraph()
        myDoc.add_paragraph()
        myDoc.add_paragraph()
        myDoc.add_paragraph()
        isFirstImg = False
    else:
        isFirstImg = True

    

myDoc.save(localDir + '\\result\\' + today + '磅单照片汇总.docx')

for fn in pictures:
    shutil.move(fn['path'], historyImgsDir)